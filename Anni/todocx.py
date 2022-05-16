import docx
import sys


# ASK ABOUT THE LANGUAGES, SHOULD ONE DOCUMENT ONLY HAVE THAT SPECIFIC LANGUAGE?
# as of now the script does that but then there is most likely gonna be a file that has less characters
# if we want to min max and mixmatch languages in the files then we just need to cat everything at once
# and not give a language as an argument

# this script makes 26 docx files, the last ones have less characters
# => max should be 25 and they should each have about 1 million characters
# => probably just delete the swe 65k char file?

# and try to downsample so that eng and fi are maxed out (to get better result should use fractions of percents)
# the en file had 714k and fi had 843k
# now fi has 912k and en 926k (this relies on the last swe file being deleted and every language being in their own docx)





# labels will be in corresponding files in their own folder

# no need for ids if we just check that the text is red and/or bolded

# def create_ids(amount):
#     ids = []
#     for i in range(amount):
#         # some great way to create a new id that does not get destroyed by deepL
#         # numbers will most likely get translated and some words might too
#         id = "id" #+ str(i)
#         ids.append(id) 

#     return ids


def make_docx(data):
    lang = sys.argv[1]
    #format the data, all should be in the same way and there should be no faulty lines
    for i in range(len(data)):
        data[i] = data[i].replace("\n", "")
        data[i] = data[i].split("\t")
    doc = docx.Document()
    doc_count = 0 # get the number for the file (100 files)
    char_count = 0 #total length of documents shouldn't be bigger than 1 million chars (maybe a bit less)

    labels = []
    texts= [one[1] for one in data]

    # here we should make an id for everything
    # and add them to both the docx and label file
    # ids = create_ids(len(texts))

    # # save the labels by printing them
    # for i in range(len(labels)):
    #     print(labels[i])

# => for loop for the array (current file)
#  check how long the the line(=string) is with len(line) and add it to char_count
    for i in range(len(texts)):
        #print(len(line))
        # check whether the document has the right amount of characters
        if char_count + len(texts[i]) + len("id") + len("\n") > 1000000: # what should be the exact number to make sure everything fits but does not go over?
            doc.save(f"../ForDeepL/{lang}_{doc_count:03d}.docx")
            
            # save the current labels as well to their own file
            with open(f"../ForDeepL/labels/{lang}_{doc_count:03d}.labels.txt", 'w') as f:
                # final = []
                # for i in range(len(labels)):
                #     dummy = labels[i][0] # + '\t' + labels[i][1]
                #     final.append(dummy)
                
                # for i in range(len(final)):
                #     #string = final[i]
                #     #f.write('%s\n' % string)
                #     f.write(final[i]+ '%s\n')
                for item in labels:
                    f.write("%s\n" % item)
            labels = []
            doc_count += 1
            #reset the file and the char count
            char_count = 0
            doc = docx.Document()
        else:
            # append the label for the text that goes into that doc
            # line = ["ids[i]", data[i][0]]
            labels.append(data[i][0])

            char_count = char_count + len(texts[i]) + len("id") + len("\n")

            # make id that is colored and bolded in the document
            parapgraph = doc.add_paragraph()
            run = parapgraph.add_run("id")
            run.bold = True
            font = run.font # bool object has no attribute font?
            font.color.rgb = docx.shared.RGBColor(255, 0, 0) # color red

            # # make a new line (new empty paragraph)
            # doc.add_paragraph("")
            parapgraph.add_run("\n")

            # check that the characters are valid because en data has something invalid
            cleaned_string = ''.join(c for c in texts[i] if valid_xml_char_ordinal(c))

            # then we add the text to the doc
            #p = doc.add_paragraph()
            parapgraph.add_run(cleaned_string)


    # the loop ends but there is not enough characters, so we just save?        
    doc.save(f"../ForDeepL/{lang}_{doc_count:03d}.docx")
    # save the current labels as well to their own file
    with open(f'../ForDeepL/labels/{lang}_{doc_count:03d}.labels.txt', 'w') as f:
        # final = []
        # for i in range(len(labels)):
        #     dummy = labels[i][0] # + '\t' + labels[i][1]
        #     final.append(dummy)
   
        # for i in range(len(final)):
        #     string = final[i]
        #     f.write('%s\n' % string)
        for item in labels:
            f.write("%s\n" % item)

    labels = []

def valid_xml_char_ordinal(c):
    codepoint = ord(c)
    # conditions ordered by presumed frequency
    return (
        0x20 <= codepoint <= 0xD7FF or
        codepoint in (0x9, 0xA, 0xD) or
        0xE000 <= codepoint <= 0xFFFD or
        0x10000 <= codepoint <= 0x10FFFF
        )

data = sys.stdin.readlines()
make_docx(data)