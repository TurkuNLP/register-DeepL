import sys
from typing import Counter
import docx



# get the file name from the command line and save here
file_name = sys.stdin.readlines()
text = file_name[0]
new = text.replace("\n", "") 

document = docx.Document(new)

# get the labels from a tsv file
fname = sys.argv[1] # use argv to get the file 
fname = "./labels/" +  fname+ '.labels.txt'


with open (fname) as f:
    labels = f.readlines()

for i in range(len(labels)):
    labels[i] = labels[i].replace("\n", "")
    # if there are no ids then no need to split
    # labels[i] = labels[i].split("\t") # split at the labels
#print(labels)

current_id = ""
translations = [] # array where we save the translated texts with their labels
final = [] # the final final array that holds everything in tsv suitable format
total = 0 # count the characters in the docx file

for i in range(len(document.paragraphs)):
    for j in range(len(document.paragraphs[i].runs)):
        if document.paragraphs[i].runs[j].font.color.rgb == docx.shared.RGBColor(255, 0, 0) or document.paragraphs[i].runs[j].font.bold == True:
            #print("red found")
            total = total + len(document.paragraphs[i].runs[j].text)
            # do I have to do anything? if I just check that it's red
        elif document.paragraphs[i].runs[j].font.color.rgb == None or document.paragraphs[i].runs[j].font.bold == False:
            if document.paragraphs[i].runs[j].text == "\n":
                total = total + len(document.paragraphs[i].runs[j].text)
            else:
                total = total + len(document.paragraphs[i].runs[j].text)
                label = labels[i]
                text=document.paragraphs[i].runs[j].text
                line = [label, text]
                translations.append(line)
            

# # look at all the paragraphs to find the colored ids
# for paragraph in document.paragraphs: 
#     for run in paragraph.runs:
#         if run.font.color.rgb == docx.shared.RGBColor(255, 0, 0) or run.font.bold == True:
#             #print(run.text)
#             current_id = run.text
#             # it finds the id!!!
#         elif run.font.color.rgb == None or run.font.bold == False:
            # if run.text == "\n":
            #     continue
            # else:
#             #print(run.text)
#             text = run.text
#             # it also finds the text!
            
#             # find the label with the same id
#             label = ""
#             for i in range(len(labels)):
#                 if labels[i][0] == current_id:
#                     label = labels[i][1]
#                     break

#             line = [label, text]
#             translations.append(line)

           
#print(translations[:2])

# and now just turn back to tsv format and save by printing?
for i in range(len(translations)):
        dummy = translations[i][0] + '\t' + translations[i][1]
        final.append(dummy)
   
for i in range(len(final)):
    print(final[i])

#print(total) # corresponds to the number that libreoffice gives
