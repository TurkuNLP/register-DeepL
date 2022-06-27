import docx
import sys

# labels will be in corresponding files in their own folder

def create_ids(amount):
    ids = []
    for i in range(amount):
        # the id numbers will get translated in some languages but they are not that important
        # only the placing of a red id marker matters
        id = "id " + str(i)
        ids.append(id) 

    return ids


def make_docx(data):
    lang = sys.argv[1]
    for i in range(len(data)):
        data[i] = data[i].replace("\n", "")
        data[i] = data[i].split("\t")
    doc = docx.Document()
    doc_count = 0 # get the number for the file (100 files)
    char_count = 0 #total length of documents shouldn't be bigger than 1 million chars (maybe a bit less)

    labels = []
    texts= [one[1] for one in data]

    # here we should make an id for everything
    # and add them to both the docx and label file
    ids = create_ids(len(texts))
    id_count = 0 # this helps reset the id's for every docx

    # for loop for the array (current file)
    for i in range(len(texts)):
        #print(len(line))
        # check whether the document has the right amount of characters
        if char_count + len(texts[i]) + len(ids[id_count]) + len("\n") > 950000: # what should be the exact number to make sure everything fits but does not go over?
            doc.save(f"../preprocessing/ForDeepL/lol/{lang}_{doc_count:03d}.docx")
            
            # save the current labels as well to their own file
            with open(f"../preprocessing/ForDeepL/full_labels/{lang}_{doc_count:03d}.labels.txt", 'w') as f:
                # if there are the ids use this:
                final = []
                for i in range(len(labels)):
                    dummy = labels[i][0] + '\t' + labels[i][1]
                    final.append(dummy)
                for i in range(len(final)):
                    string = final[i]
                    f.write('%s\n' % string)

            labels = []
            doc_count += 1
            #reset the file and the char count
            char_count = 0
            id_count = 0
            doc = docx.Document()


            # add the current line as well so we don't skip it!!!! IMPORTANT SO NO LINES GO MISSING WHEN MAKING THE DOCX :D


            # append the id to the labels
            line = [ids[id_count], data[i][0]]
            labels.append(line)

            char_count = char_count + len(texts[i]) + len(ids[id_count]) + len("\n")

            # make id that is colored and bolded in the document
            parapgraph = doc.add_paragraph()
            run = parapgraph.add_run(ids[id_count]) # add the id (could be ids[i])
            run.bold = True
            font = run.font
            font.color.rgb = docx.shared.RGBColor(255, 0, 0) # color red

            # check that the characters are valid because en data has something invalid
            cleaned_string = ''.join(c for c in texts[i] if valid_xml_char_ordinal(c))

            # then we add the text to the doc
            p = doc.add_paragraph()
            run = p.add_run(cleaned_string)
            run.bold = False
            run.font.color.rgb = None
            
            id_count = id_count + 1

        else:
            # append the id to the labels
            line = [ids[id_count], data[i][0]]
            labels.append(line)

            char_count = char_count + len(texts[i]) + len(ids[id_count]) + len("\n")

            # make id that is colored and bolded in the document
            parapgraph = doc.add_paragraph()
            run = parapgraph.add_run(ids[id_count]) # add the id (could be ids[i])
            run.bold = True
            font = run.font
            font.color.rgb = docx.shared.RGBColor(255, 0, 0) # color red

            # check that the characters are valid because en data has something invalid
            cleaned_string = ''.join(c for c in texts[i] if valid_xml_char_ordinal(c))

            # then we add the text to the doc
            p = doc.add_paragraph()
            run = p.add_run(cleaned_string)
            run.bold = False
            run.font.color.rgb = None
            
            id_count = id_count + 1


    # the loop ends but there is not enough characters so just save       
    doc.save(f"../preprocessing/ForDeepL/lol/{lang}_{doc_count:03d}.docx")
    # save the current labels as well to their own file
    with open(f'../preprocessing/ForDeepL/full_labels/{lang}_{doc_count:03d}.labels.txt', 'w') as f:
        #if there are the ids use this:
        final = []
        for i in range(len(labels)):
            dummy = labels[i][0] + '\t' + labels[i][1]
            final.append(dummy)
        for i in range(len(final)):
            string = final[i]
            f.write('%s\n' % string)

# this is for checking that there are no weird characters that break the docx making
def valid_xml_char_ordinal(c):
    codepoint = ord(c)
    # conditions ordered by presumed frequency
    return (
        0x20 <= codepoint <= 0xD7FF or
        codepoint in (0x9, 0xA, 0xD) or
        0xE000 <= codepoint <= 0xFFFD or
        0x10000 <= codepoint <= 0x10FFFF
        )

data = sys.stdin.readlines()
make_docx(data)