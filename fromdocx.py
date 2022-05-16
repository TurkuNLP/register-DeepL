import sys
from typing import Counter
import docx


# I still somehow have to check that it has not skipped a label?


# get the file name from the command line and save here
file_name = sys.stdin.readlines()
text = file_name[0]
new = text.replace("\n", "") 

document = docx.Document(new)

# get the labels from a tsv file
fname = sys.argv[1] # use argv to get the file 
fname = "./labels/" +  fname+ '.labels.txt'


with open (fname) as f:
    labels = f.readlines()

for i in range(len(labels)):
    labels[i] = labels[i].replace("\n", "")
    # if there are no ids then no need to split
    labels[i] = labels[i].split("\t")
#print(labels)

current_id = ""
translations = [] # array where we save the translated texts with their labels
final = [] # the final final array that holds everything in tsv suitable format
total = 0 # count the characters in the docx file
id_found = False
count = 0

for i in range(len(document.paragraphs)):
    for j in range(len(document.paragraphs[i].runs)):
        if document.paragraphs[i].runs[j].font.color.rgb == docx.shared.RGBColor(255, 0, 0) or document.paragraphs[i].runs[j].font.bold == True:
            total = total + len(document.paragraphs[i].runs[j].text)
            current_id = document.paragraphs[i].runs[j].text
            #print(current_id)
            # do I have to do anything? if I just check that it's red
            id_found = True
        elif document.paragraphs[i].runs[j].font.color.rgb == None and document.paragraphs[i].runs[j].font.bold == False:
            # if document.paragraphs[i].runs[j].text == "\n":
            #     total = total + len(document.paragraphs[i].runs[j].text)
            # else:
            if id_found == True:
                #print("text really found")
                text=document.paragraphs[i].runs[j].text
                total = total + len(document.paragraphs[i].runs[j].text)
                #label = labels[(i-1) / 2][1] 
                
                label = ""
                c = ''.join(filter(str.isdigit, current_id))
                if c == "":
                    print("The number in the id is broken", current_id)
                    count = count + 1
                    continue
                # check that no text has been skipped by looking at the text count and comparing it to the current id
                elif count != int(c):
                    # if count is smaller than id
                    count = count + (c - count)
                    continue
                    #raise Exception("A piece of text has been skipped")

                # # find the label with the same id
                # for i in range(len(labels)):
                #     if c in labels[i][0]:
                #         label = labels[i][1]
                #         break
                #  if label == "":
                #     continue 
                #     #raise Exception("the correct ID wasn't found",current_id)

                # OOR use the count to get the right label, I guess the former is more robust if the numbers are still numbers
                label = labels[count][1]
                
                line = [label, text]
                translations.append(line)
                id_found = False
                count = count +1
            else:
                continue
                #raise Exception("the id is missing", current_id)
            

# look at all the paragraphs to find the colored ids
# for paragraph in document.paragraphs: 
#     for run in paragraph.runs:
#         if run.font.color.rgb == docx.shared.RGBColor(255, 0, 0) or run.font.bold == True:
#             #print(run.text)
#             current_id = run.text
#             # it finds the id!!!
#             id_found = True
#         elif run.font.color.rgb == None and run.font.bold == False:
#             if id_found == True:
#                 #print(run.text)
#                 text = run.text
#                 # it also finds the text!
                
#                 # find the label with the same id
#                 c = ''.join(filter(str.isdigit, current_id))
#                 label = ""
#                 for i in range(len(labels)):
#                     if labels[i][0] == current_id:
#                         label = labels[i][1]
#                         break
#                 if label == "":
# #                     #continue #?
#                     raise Exception("the correct ID wasn't found",current_id)
#                 line = [label, text]
#                 translations.append(line)
#                 id_found = False
#             else:
#                 raise Exception("the correct ID wasn't found", current_id)
                

           
#print(translations[:2])

# and now just turn back to tsv format and save by printing?
for i in range(len(translations)):
        dummy = translations[i][0] + '\t' + translations[i][1]
        final.append(dummy)
   
for i in range(len(final)):
    print(final[i])

#print(total) # corresponds to the number that libreoffice gives
